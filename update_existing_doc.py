import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

def update_document():
    file_path = 'HabitFlow_Project_Documentation_Aditya_Dubey.docx'
    
    if not os.path.exists(file_path):
        print(f"Error: {file_path} not found.")
        return

    print("Opening existing document...")
    doc = Document(file_path)

    # 1. Append Acknowledgement to the end of the document
    doc.add_page_break()
    doc.add_heading('Chapter 7: Acknowledgement', level=1)
    
    ack_text = (
        "I would like to express my deepest appreciation to all those who provided me the possibility to complete this report. "
        "A special gratitude I give to our final year project manager, whose contribution in stimulating suggestions and encouragement, "
        "helped me to coordinate my project especially in writing this report.\n\n"
        "Furthermore, I would also like to acknowledge with much appreciation the crucial role of the staff, who gave the permission "
        "to use all required equipment and the necessary materials to complete the task \"HabitFlow\".\n\n"
        "Last but not least, many thanks go to the head of the project, who have invested his full effort in guiding the team in achieving the goal. "
        "I have to appreciate the guidance given by other supervisors as well as the panels especially in our project presentation that has improved "
        "our presentation skills thanks to their comment and advices."
    )
    doc.add_paragraph(ack_text)

    # 2. Insert Index before Chapter 1
    # Find the paragraph that says "Chapter 1: Introduction"
    target_p = None
    for p in doc.paragraphs:
        if 'Chapter 1: Introduction' in p.text:
            target_p = p
            break
            
    if target_p:
        print("Found Chapter 1. Inserting Index before it...")
        # Create a heading for the Index
        heading = target_p.insert_paragraph_before('INDEX')
        heading.style = 'Heading 1'
        
        # Add Index items
        items = [
            "1. Introduction & Abstract",
            "2. Functional Design Document (FDD)",
            "3. Technical Design Document (TDD)",
            "4. Environment Configuration",
            "5. Component Level Specifications",
            "6. Quality Assurance & Testing",
            "7. Conclusion",
            "8. Acknowledgement"
        ]
        
        for item in items:
            target_p.insert_paragraph_before(item)
            
        # Add a page break after the Index (before Chapter 1)
        pb_paragraph = target_p.insert_paragraph_before('')
        run = pb_paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
    else:
        print("Could not find Chapter 1. Index not inserted.")

    # Save the document
    doc.save(file_path)
    print("Successfully updated the document without modifying the first page!")

if __name__ == '__main__':
    update_document()
